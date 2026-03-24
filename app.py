import os
import io
import json
import streamlit as st
from dotenv import load_dotenv
from docx import Document
from docx.shared import Pt
from PyPDF2 import PdfReader

# Load API key
load_dotenv()
api_key = os.getenv("GEMINI_API_KEY")

# Gemini Client
from google.genai import Client
client = Client(api_key=api_key)

MODEL_NAME = "gemini-2.5-flash"

# -------------------------
# CONFIG
# -------------------------
st.set_page_config(page_title="🌍 AI Translator Agent", layout="centered")

LANGUAGES = [
    "English", "Hindi", "Spanish", "French", "German",
    "Chinese", "Japanese", "Korean", "Arabic", "Russian"
]

TONES = ["Neutral", "Formal", "Casual", "Professional", "Friendly"]

# -------------------------
# HELPERS
# -------------------------
def read_file(uploaded_file):
    if not uploaded_file:
        return ""

    name = uploaded_file.name.lower()

    if name.endswith(".txt"):
        return uploaded_file.read().decode("utf-8")

    if name.endswith(".docx"):
        doc = Document(uploaded_file)
        return "\n".join(p.text for p in doc.paragraphs)

    if name.endswith(".pdf"):
        reader = PdfReader(uploaded_file)
        return "\n".join(page.extract_text() or "" for page in reader.pages)

    return ""

def generate_translation(text, target_lang, tone):
    prompt = f"""
    You are an expert translator.

    Translate the following text into {target_lang}.
    Maintain a {tone} tone.

    Also:
    - Preserve meaning accurately
    - Improve clarity if needed
    - Keep it natural and fluent

    Text:
    {text}
    """

    try:
        res = client.models.generate_content(
            model=MODEL_NAME,
            contents=prompt
        )
        return res.text
    except:
        return "❌ Translation failed"

def generate_docx(text):
    doc = Document()
    for line in text.split("\n"):
        p = doc.add_paragraph()
        run = p.add_run(line)
        run.font.size = Pt(11)

    bio = io.BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio.read()

# -------------------------
# UI
# -------------------------
st.title("🌍 AI Language Translation Agent")

source_lang = st.selectbox("Source Language", ["Auto Detect"] + LANGUAGES)
target_lang = st.selectbox("Target Language", LANGUAGES)
tone = st.selectbox("Tone", TONES)

uploaded_file = st.file_uploader("Upload file", type=["txt", "docx", "pdf"])
file_text = read_file(uploaded_file)

user_text = st.text_area("Enter text", value=file_text, height=200)

col1, col2 = st.columns(2)
translate_btn = col1.button("🚀 Translate")
clear_btn = col2.button("🧹 Clear")

if clear_btn:
    st.session_state.output = ""

# -------------------------
# TRANSLATION
# -------------------------
if translate_btn:
    if not user_text.strip():
        st.warning("Please enter text")
        st.stop()

    with st.spinner("Translating..."):
        output = generate_translation(user_text, target_lang, tone)
        st.session_state.output = output

# -------------------------
# OUTPUT
# -------------------------
if "output" in st.session_state and st.session_state.output:
    result = st.session_state.output

    # TTS (browser)
    escaped = json.dumps(result)
    tts_html = f"""
    <button onclick="speak()">🔊 Speak</button>
    <script>
    function speak(){{
        let msg = new SpeechSynthesisUtterance({escaped});
        window.speechSynthesis.speak(msg);
    }}
    </script>
    """
    st.markdown(tts_html, unsafe_allow_html=True)

    st.subheader("📄 Translated Text")
    st.text_area("", result, height=200)

    # Stats
    words = len(result.split())
    st.caption(f"Words: {words} | Read time: {words//200 + 1} min")

    # Downloads
    st.download_button("⬇ Download TXT", result, "translation.txt")
    st.download_button("⬇ Download DOCX", generate_docx(result), "translation.docx")

    # Improve Section
    with st.expander("✨ Improve Translation"):
        improve_prompt = f"Improve this translation:\n{result}"
        res = client.models.generate_content(
            model=MODEL_NAME,
            contents=improve_prompt
        )
        st.write(res.text)